"""
Read the first table from a simple Word (.docx) document into a DataFrame.

This is intentionally minimal: it reads only the first table, treats its first
row as the header, and tries to convert each column to numbers where possible
(Word cells are always text, but our models expect numeric features).
"""

import pandas as pd


def read_word_table(uploaded_file):
    """
    Convert the first table in a .docx file to a pandas DataFrame.

    Parameters
    ----------
    uploaded_file : file-like object
        An uploaded .docx file.

    Returns
    -------
    pandas.DataFrame

    Raises
    ------
    ValueError
        If the document contains no table.
    """
    # Imported here so python-docx is only required when a Word file is used.
    from docx import Document

    document = Document(uploaded_file)

    if not document.tables:
        raise ValueError(
            "No table found in the Word document. Please upload CSV or Excel."
        )

    table = document.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    header, *data_rows = rows
    df = pd.DataFrame(data_rows, columns=header)

    # Word stores everything as text, so coerce numeric-looking columns to
    # numbers. Columns that are genuinely categorical (e.g. job, purpose)
    # stay as text because the conversion leaves them unchanged on failure.
    for col in df.columns:
        converted = pd.to_numeric(df[col], errors="coerce")
        if converted.notna().all():
            df[col] = converted

    return df
